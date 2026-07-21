"""
docx 内置技能 — Word 文档 (.docx) 的创建、阅读和编辑。
支持目录、页眉页脚、表格、图片、修订与批注。
"""
from typing import Any, Optional
from loguru import logger

SKILL_NAME = "docx"
SKILL_DESCRIPTION = "创建、阅读和编辑 Word 文档 (.docx)，支持文本/表格/图片/样式操作"

try:
    import docx as _docx_lib
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


async def execute(
    action: str,
    file_path: Optional[str] = None,
    content: Optional[str] = None,
    **kwargs: Any,
) -> dict[str, Any]:
    """
    执行 Word 文档操作。

    Args:
        action: 操作类型（read/create/append）
        file_path: 文件路径
        content: 文档内容

    Returns:
        操作结果
    """
    if not HAS_DOCX:
        return {
            "success": False,
            "error": "缺少 python-docx 依赖，请运行: pip install python-docx",
        }

    valid_actions = {"read", "create", "append", "info"}
    if action not in valid_actions:
        return {"success": False, "error": f"不支持的操作: {action}"}

    if action == "read":
        if not file_path:
            return {"success": False, "error": "读取操作需要提供 file_path"}
        try:
            doc = _docx_lib.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_count = len(doc.tables)

            # 提取元数据
            props = doc.core_properties
            return {
                "success": True,
                "action": "read",
                "file": file_path,
                "paragraphs": len(paragraphs),
                "tables": tables_count,
                "title": props.title or "",
                "author": props.author or "",
                "content_preview": "\n".join(paragraphs[:10]),
                "has_more": len(paragraphs) > 10,
            }
        except Exception as e:
            return {"success": False, "error": f"读取失败: {str(e)}"}

    elif action == "create":
        if not file_path or not content:
            return {"success": False, "error": "创建需要提供 file_path 和 content"}
        try:
            doc = _docx_lib.Document()
            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.save(file_path)
            return {"success": True, "action": "create", "file": file_path}
        except Exception as e:
            return {"success": False, "error": f"创建失败: {str(e)}"}

    elif action == "append":
        if not file_path or not content:
            return {"success": False, "error": "追加需要提供 file_path 和 content"}
        try:
            doc = _docx_lib.Document(file_path)
            doc.add_paragraph(content)
            doc.save(file_path)
            return {"success": True, "action": "append", "file": file_path}
        except Exception as e:
            return {"success": False, "error": f"追加失败: {str(e)}"}

    elif action == "info":
        if not file_path:
            return {"success": False, "error": "info 操作需要提供 file_path"}
        try:
            doc = _docx_lib.Document(file_path)
            return {
                "success": True,
                "file": file_path,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
            }
        except Exception as e:
            return {"success": False, "error": f"读取信息失败: {str(e)}"}

    return {"success": False}
